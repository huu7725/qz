from fastapi import FastAPI, UploadFile, File, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import json, random, re, os, base64, io, traceback
from datetime import datetime
from pathlib import Path

# ══ CẤU HÌNH - THAY KEY VÀO ĐÂY ══
ANTHROPIC_API_KEY = "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"
# ══════════════════════════════════

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

DATA_DIR = Path("data/users")
DATA_DIR.mkdir(parents=True, exist_ok=True)
CONFIG_FILE = Path("data/config.json")
SESSIONS: dict = {}

def user_dir(uid):
    d = DATA_DIR / re.sub(r'[^\w]', '_', uid or "guest")
    d.mkdir(parents=True, exist_ok=True)
    return d

def files_index_path(uid): return user_dir(uid) / "files_index.json"
def history_path(uid):     return user_dir(uid) / "history.json"

def load_json(path, default):
    try:
        if Path(path).exists():
            with open(path, encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return default

def save_json(path, data):
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def get_uid(x_user_id: str = "") -> str:
    return (x_user_id or "guest").strip()

def load_config():
    cfg = {"anthropic_key": ANTHROPIC_API_KEY, "ai_parse_enabled": True}
    try:
        if CONFIG_FILE.exists():
            saved = load_json(CONFIG_FILE, {})
            cfg["ai_parse_enabled"] = saved.get("ai_parse_enabled", True)
    except:
        pass
    return cfg

def has_valid_key():
    k = ANTHROPIC_API_KEY.strip()
    return bool(k and k != "YOUR_KEY_HERE" and k.startswith("sk-ant-"))

# ══════════════════════════════════════════
#  PDF PARSER — BẢNG TÔ MÀU NỀN
# ══════════════════════════════════════════

def is_highlight(fill):
    """Nhận dạng màu highlight - bỏ trắng/đen"""
    if not fill or len(fill) < 3:
        return False
    r, g, b = float(fill[0]), float(fill[1]), float(fill[2])
    if r > 0.92 and g > 0.92 and b > 0.92: return False  # trắng
    if r < 0.05 and g < 0.05 and b < 0.05: return False  # đen
    if r > 0.85 and g > 0.85 and b > 0.85: return False  # gần trắng
    # Cyan/teal (file pháp luật của bạn)
    if g > 0.50 and b > 0.40 and r < 0.60: return True
    # Xanh lá
    if g > 0.55 and g > r * 1.3 and g > b * 1.2: return True
    # Vàng
    if r > 0.75 and g > 0.70 and b < 0.40: return True
    # Hồng/cam
    if r > 0.80 and g < 0.65 and b < 0.65: return True
    # Xanh dương
    if b > 0.55 and b > r * 1.2 and g > 0.45: return True
    return False

def get_highlight_rects(page):
    """Lấy tất cả rect tô màu highlight"""
    rects = []
    try:
        for d in page.get_drawings():
            if is_highlight(d.get("fill")):
                rects.append(d["rect"])
    except Exception as e:
        print(f"get_drawings error: {e}")
    return rects

def span_in_rect(bbox, rects, pad=4):
    cx = (bbox[0] + bbox[2]) / 2
    cy = (bbox[1] + bbox[3]) / 2
    for r in rects:
        if r.x0-pad <= cx <= r.x1+pad and r.y0-pad <= cy <= r.y1+pad:
            return True
    return False

def get_all_spans(page):
    """Lấy tất cả text spans với toạ độ"""
    spans = []
    try:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                for sp in line["spans"]:
                    t = sp["text"].strip()
                    if t:
                        bbox = sp["bbox"]
                        spans.append({
                            "text": t,
                            "bbox": bbox,
                            "cx": (bbox[0]+bbox[2])/2,
                            "cy": (bbox[1]+bbox[3])/2,
                        })
    except Exception as e:
        print(f"get_all_spans error: {e}")
    return spans

def group_rows(spans, tol=12):
    """Nhóm spans theo hàng ngang (y gần nhau)"""
    rows = []
    for sp in sorted(spans, key=lambda s: s["cy"]):
        placed = False
        for row in rows:
            if abs(row["cy"] - sp["cy"]) < tol:
                row["spans"].append(sp)
                row["cy"] = (row["cy"] + sp["cy"]) / 2
                placed = True
                break
        if not placed:
            rows.append({"cy": sp["cy"], "spans": [sp]})
    return sorted(rows, key=lambda r: r["cy"])

def parse_pdf_table(page, hl_rects):
    """
    Parse PDF dạng BẢNG:
    TT | DVKT | Nội dung câu hỏi | Đáp án A | Đáp án B | Đáp án C | Đáp án D
    Ô nào tô màu = đáp án đúng của hàng đó
    """
    spans = get_all_spans(page)
    if not spans:
        return None

    rows = group_rows(spans)

    # 1. Tìm header row chứa cột đáp án
    col_A = col_B = col_C = col_D = col_Q = None
    header_y = None

    for row in rows:
        full_text = " ".join(s["text"] for s in row["spans"]).upper().replace(" ", "")
        has_da = bool(re.search(r'ĐÁPÁN|ANSWER', " ".join(s["text"] for s in row["spans"]), re.IGNORECASE))
        labels = [s for s in row["spans"] if s["text"].strip().upper() in ["A", "B", "C", "D"]]

        if has_da or len(labels) >= 3:
            header_y = row["cy"]
            for sp in row["spans"]:
                t = sp["text"].upper().replace(" ", "")
                if any(k in t for k in ["NỘIDUNG", "CÂUHỎI", "CONTENT", "NOIUNG"]):
                    col_Q = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]NA|ANSWERA', t) or t == "A":
                    col_A = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]NB|ANSWERB', t) or t == "B":
                    col_B = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]NC|ANSWERC', t) or t == "C":
                    col_C = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]ND|ANSWERD', t) or t == "D":
                    col_D = sp["cx"]
            break

    # Nếu không tìm đủ cột, trả None
    found_cols = sum(1 for c in [col_A, col_B, col_C, col_D] if c is not None)
    if found_cols < 2:
        return None

    col_map = {}
    for label, cx in [("A", col_A), ("B", col_B), ("C", col_C), ("D", col_D)]:
        if cx is not None:
            col_map[label] = cx

    # 2. Map hl_rects → (row_num, label)
    answer_map = {}
    for hl in hl_rects:
        hcx = (hl.x0 + hl.x1) / 2
        hcy = (hl.y0 + hl.y1) / 2
        if header_y and hcy <= header_y + 5:
            continue

        # Cột nào gần nhất
        best_label, best_dist = None, 999
        for label, cx in col_map.items():
            d = abs(hcx - cx)
            if d < best_dist:
                best_dist = d
                best_label = label
        if not best_label or best_dist > 120:
            continue

        # Hàng nào? Tìm số TT gần nhất theo y
        best_tt, best_dy = None, 999
        for sp in spans:
            if re.match(r'^\d+$', sp["text"]):
                dy = abs(sp["cy"] - hcy)
                if dy < best_dy:
                    best_dy = dy
                    best_tt = int(sp["text"])
        if best_tt and best_dy < 30:
            answer_map[best_tt] = best_label

    if not answer_map:
        return None

    # 3. Ghép câu hỏi từ các hàng
    questions = []
    data_rows = [r for r in rows if header_y is None or r["cy"] > header_y + 8]

    # Gom nhiều row liên tiếp cho cùng 1 câu (câu dài nhiều dòng)
    tt_to_spans = {}
    current_tt = None
    for row in data_rows:
        # Có số TT không?
        tt_spans = [s for s in row["spans"] if re.match(r'^\d+$', s["text"])]
        if tt_spans:
            current_tt = int(tt_spans[0]["text"])
        if current_tt:
            if current_tt not in tt_to_spans:
                tt_to_spans[current_tt] = []
            tt_to_spans[current_tt].extend(row["spans"])

    for tt, row_spans in tt_to_spans.items():
        if tt not in answer_map:
            continue

        # Nội dung câu hỏi: spans ở cột "Nội dung" (hoặc giữa bảng)
        if col_Q:
            q_spans = [s for s in row_spans
                       if col_Q - 180 <= s["cx"] <= col_Q + 180
                       and s["text"] not in [str(tt)]]
        else:
            # Lấy spans không phải số đơn và không phải label đơn
            q_spans = [s for s in row_spans
                       if len(s["text"]) > 2
                       and not re.match(r'^[\dABCD]$', s["text"])
                       and not any(abs(s["cx"] - cx) < 80 for cx in col_map.values())]

        q_text = " ".join(s["text"] for s in sorted(q_spans, key=lambda x: (x["cy"], x["cx"])))
        if not q_text.strip():
            continue

        # Đáp án từng cột
        choices = []
        for label in sorted(col_map.keys()):
            cx = col_map[label]
            c_spans = [s for s in row_spans
                       if abs(s["cx"] - cx) < 90 and s["text"] not in [str(tt)]]
            c_text = " ".join(s["text"] for s in sorted(c_spans, key=lambda x: (x["cy"], x["cx"])))
            if c_text.strip():
                choices.append({"label": label, "text": c_text.strip()})

        if len(choices) >= 2:
            questions.append({
                "id": tt,
                "question": q_text.strip(),
                "choices": choices,
                "answer": answer_map.get(tt, ""),
                "explanation": ""
            })

    return questions if questions else None

def parse_pdf_inline(page, hl_rects):
    """Parse PDF dạng inline: Câu 1: ... A. ... B. ... C. ... D."""
    lines = []
    try:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                line_text, hl = "", False
                for sp in line["spans"]:
                    line_text += sp["text"]
                    if hl_rects and span_in_rect(sp["bbox"], hl_rects):
                        hl = True
                t = line_text.strip()
                if t:
                    lines.append({"text": t, "hl": hl})
    except Exception as e:
        print(f"parse_pdf_inline error: {e}")
        return []

    questions, cur = [], None
    for item in lines:
        text, hl = item["text"], item["hl"]
        m = re.match(r'^(?:C[âa]u\s*)?(\d+)[\.\:\)]\s*(.+)', text, re.IGNORECASE)
        if m and not re.match(r'^[A-Da-d][\.\)]', m.group(2)):
            if cur and len(cur["choices"]) >= 2:
                questions.append(cur)
            cur = {"id": int(m.group(1)), "question": m.group(2).strip(),
                   "choices": [], "answer": "", "explanation": ""}
        elif re.match(r'^[A-Da-d][\.\)]\s+.+', text) and cur:
            label = text[0].upper()
            cur["choices"].append({
                "label": label,
                "text": re.sub(r'^[A-Da-d][\.\)]\s*', '', text).strip()
            })
            if hl and not cur["answer"]:
                cur["answer"] = label
        elif re.match(r'^(Đ[áa]p\s*[áa]n|ĐA)\s*[:\-]', text, re.IGNORECASE) and cur:
            mm = re.search(r'[A-D]', text.upper())
            if mm and not cur["answer"]:
                cur["answer"] = mm.group()
    if cur and len(cur["choices"]) >= 2:
        questions.append(cur)
    return questions

def parse_pdf(content: bytes) -> list:
    """Parser PDF chính: thử bảng → fallback inline"""
    import fitz
    doc = fitz.open(stream=content, filetype="pdf")
    all_questions = []

    for page in doc:
        hl_rects = get_highlight_rects(page)

        # Thử parse dạng bảng trước
        if hl_rects:
            try:
                table_qs = parse_pdf_table(page, hl_rects)
                if table_qs:
                    all_questions.extend(table_qs)
                    continue
            except Exception as e:
                print(f"Table parse failed: {e}")

        # Fallback dạng inline
        try:
            inline_qs = parse_pdf_inline(page, hl_rects)
            all_questions.extend(inline_qs)
        except Exception as e:
            print(f"Inline parse failed: {e}")

    # Re-index
    for i, q in enumerate(all_questions):
        q["id"] = i + 1
    return all_questions

# ══════════════════════════════════════════
#  WORD PARSER
# ══════════════════════════════════════════

def parse_word(content: bytes) -> list:
    import docx
    doc = docx.Document(io.BytesIO(content))
    answer_map = {}
    for table in doc.tables:
        cells = [c.text.strip() for row in table.rows for c in row.cells]
        is_ans = any(re.search(r'ĐÁP\s*ÁN|ĐA\b', x, re.IGNORECASE) for x in cells)
        if not is_ans:
            is_ans = sum(1 for x in cells if x.upper() in 'ABCD') > len(cells) * 0.3
        if is_ans:
            for row in table.rows:
                vals = [c.text.strip() for c in row.cells]
                for i in range(0, len(vals)-1, 2):
                    if re.match(r'^\d+$', vals[i]) and vals[i+1].upper() in 'ABCD':
                        answer_map[int(vals[i])] = vals[i+1].upper()
            rows = list(table.rows)
            if len(rows) >= 2:
                header = [c.text.strip() for c in rows[0].cells]
                for row in rows[1:]:
                    for j, c in enumerate(row.cells):
                        v = c.text.strip().upper()
                        if j < len(header) and re.match(r'^\d+$', header[j]) and v in 'ABCD':
                            answer_map[int(header[j])] = v
            if answer_map:
                break
    questions = _parse_lines([p.text for p in doc.paragraphs])
    for q in questions:
        if not q["answer"] and q["id"] in answer_map:
            q["answer"] = answer_map[q["id"]]
    return questions

def _parse_lines(lines):
    questions, cur = [], None
    for line in lines:
        line = line.strip() if isinstance(line, str) else ""
        if not line:
            continue
        m = re.match(r'^(?:C[âa]u\s*)?(\d+)[\.\:\)]\s*(.+)', line, re.IGNORECASE)
        if m and not re.match(r'^[A-Da-d][\.\)]', m.group(2)):
            if cur and len(cur["choices"]) >= 2:
                questions.append(cur)
            cur = {"id": int(m.group(1)), "question": m.group(2).strip(),
                   "choices": [], "answer": "", "explanation": ""}
        elif re.match(r'^[A-Da-d][\.\)]\s+.+', line) and cur:
            label = line[0].upper()
            cur["choices"].append({
                "label": label,
                "text": re.sub(r'^[A-Da-d][\.\)]\s*', '', line).strip()
            })
        elif re.match(r'^(Đ[áa]p\s*[áa]n|ĐA)', line, re.IGNORECASE) and cur:
            mm = re.search(r'[A-D]', line.upper())
            if mm:
                cur["answer"] = mm.group()
    if cur and len(cur["choices"]) >= 2:
        questions.append(cur)
    return questions

# ══════════════════════════════════════════
#  CLAUDE VISION AI
# ══════════════════════════════════════════

VISION_PROMPT = """Đây là ảnh trang đề thi trắc nghiệm tiếng Việt.
Đáp án ĐÚNG = ô được TÔ MÀU (cyan, xanh, vàng, hồng...) trong cột đáp án.
Nếu là bảng (TT | Nội dung | Đáp án A | B | C | D): cột tô màu = đáp án đúng của hàng đó.
Trả về JSON THUẦN TÚY không có markdown:
[{"id":1,"question":"...","choices":[{"label":"A","text":"..."},{"label":"B","text":"..."},{"label":"C","text":"..."},{"label":"D","text":"..."}],"answer":"A"}]
answer="" nếu không xác định. Không bỏ sót câu."""

def parse_with_vision_ai(content: bytes, filetype: str, api_key: str) -> list:
    try:
        import anthropic
        from pdf2image import convert_from_bytes
    except ImportError:
        raise HTTPException(500, "Cần: pip install anthropic pdf2image")

    client = anthropic.Anthropic(api_key=api_key)
    if filetype != "pdf":
        return parse_word(content)

    try:
        images = convert_from_bytes(content, dpi=150)
    except Exception as e:
        raise HTTPException(500, f"Lỗi convert PDF→ảnh: {e}")

    all_qs = []
    for i, img in enumerate(images):
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=90)
        img_b64 = base64.b64encode(buf.getvalue()).decode()
        try:
            resp = client.messages.create(
                model="claude-opus-4-5", max_tokens=4096,
                messages=[{"role": "user", "content": [
                    {"type": "image", "source": {
                        "type": "base64", "media_type": "image/jpeg", "data": img_b64}},
                    {"type": "text", "text": VISION_PROMPT}
                ]}]
            )
            raw = resp.content[0].text.strip()
            m = re.search(r'\[.*\]', raw, re.DOTALL)
            if m:
                page_qs = json.loads(m.group())
                for q in page_qs:
                    q["id"] = len(all_qs) + page_qs.index(q) + 1
                all_qs.extend(page_qs)
        except Exception as e:
            print(f"Vision AI p{i+1}: {e}")
    return all_qs

def smart_parse(content: bytes, filename: str, force_ai: bool = False) -> dict:
    fname_lower = filename.lower()
    filetype = "pdf" if fname_lower.endswith(".pdf") else "docx"
    cfg = load_config()
    ai_key = cfg.get("anthropic_key", "").strip()
    ai_ok = bool(ai_key and ai_key != "YOUR_KEY_HERE")
    ai_enabled = cfg.get("ai_parse_enabled", True) and ai_ok

    method, error_msg = "normal", ""
    try:
        questions = parse_pdf(content) if filetype == "pdf" else parse_word(content)
    except Exception as e:
        questions = []
        error_msg = str(e)
        print(f"Parse error:\n{traceback.format_exc()}")

    total = len(questions)
    has_ans = sum(1 for q in questions if q.get("answer") in list("ABCD"))
    ans_rate = round(has_ans / total * 100) if total > 0 else 0

    need_ai = force_ai or (ai_enabled and ((total > 0 and ans_rate < 30) or total == 0))
    if need_ai and ai_ok:
        method = "vision_ai"
        try:
            questions = parse_with_vision_ai(content, filetype, ai_key)
            total = len(questions)
            has_ans = sum(1 for q in questions if q.get("answer") in list("ABCD"))
            ans_rate = round(has_ans / total * 100) if total > 0 else 0
        except Exception as e:
            print(f"Vision AI error: {e}")

    return {"questions": questions, "method": method, "total": total,
            "has_ans": has_ans, "ans_rate": ans_rate,
            "ai_available": ai_ok, "error": error_msg}

# ══════════════════════════════════════════
#  API ENDPOINTS
# ══════════════════════════════════════════

@app.get("/config")
def get_config():
    ok = has_valid_key()
    cfg = load_config()
    return {"ai_parse_enabled": cfg.get("ai_parse_enabled", True),
            "has_api_key": ok, "key_preview": "sk-ant-...***" if ok else ""}

class ConfigBody(BaseModel):
    anthropic_key: str = ""
    ai_parse_enabled: bool = True

@app.post("/config")
def update_config(body: ConfigBody):
    save_json(CONFIG_FILE, {"ai_parse_enabled": body.ai_parse_enabled})
    return {"message": "Đã lưu cài đặt", "has_key": has_valid_key()}

@app.delete("/config/key")
def delete_api_key():
    save_json(CONFIG_FILE, {"ai_parse_enabled": True})
    return {"message": "Đã reset config"}

@app.get("/stats")
def get_stats(x_user_id: str = Header(default="guest")):
    uid = get_uid(x_user_id)
    index = load_json(files_index_path(uid), {})
    history = load_json(history_path(uid), [])
    ok = has_valid_key()
    cfg = load_config()
    return {
        "total_questions": sum(f["count"] for f in index.values()),
        "with_answer": sum(f["with_answer"] for f in index.values()),
        "total_sessions": len(history),
        "avg_score": round(sum(h["percent"] for h in history)/len(history)) if history else 0,
        "best_score": max((h["percent"] for h in history), default=0),
        "ai_available": ok,
        "ai_enabled": cfg.get("ai_parse_enabled", True),
        "files": [
            {"id": fid, "name": f["name"], "count": f["count"],
             "with_answer": f["with_answer"], "uploaded_at": f["uploaded_at"],
             "parse_method": f.get("parse_method", "normal")}
            for fid, f in index.items()
        ]
    }

@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    x_user_id: str = Header(default="guest"),
    x_force_ai: str = Header(default="false")
):
    uid = get_uid(x_user_id)
    content = await file.read()
    fname = file.filename or "file"

    if not (fname.lower().endswith(".docx") or fname.lower().endswith(".pdf")):
        raise HTTPException(400, "Chỉ hỗ trợ .docx và .pdf")

    force_ai = x_force_ai.lower() == "true"
    try:
        result = smart_parse(content, fname, force_ai=force_ai)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Lỗi xử lý: {e}")

    questions = result["questions"]
    if not questions:
        err = result.get("error", "")
        raise HTTPException(400, f"Không tìm thấy câu hỏi.{' Chi tiết: '+err if err else ' Thử bật AI Vision.'}")

    base_name = re.sub(r'\.(docx|pdf)$', '', fname, flags=re.IGNORECASE)
    file_id = re.sub(r'[^\w\-]', '_', base_name)[:50]

    index = load_json(files_index_path(uid), {})
    q_file = user_dir(uid) / f"{file_id}.json"
    existing = load_json(q_file, [])
    exist_keys = {q["question"][:60].lower() for q in existing}
    added, max_id = 0, max((q.get("id", 0) for q in existing), default=0)
    for q in questions:
        key = q["question"][:60].lower()
        if key not in exist_keys:
            max_id += 1
            q["id"] = max_id
            existing.append(q)
            exist_keys.add(key)
            added += 1

    save_json(q_file, existing)
    has_ans = sum(1 for q in existing if q.get("answer") in list("ABCD"))
    index[file_id] = {
        "name": base_name, "filename": fname,
        "count": len(existing), "with_answer": has_ans,
        "uploaded_at": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "file_id": file_id, "parse_method": result["method"]
    }
    save_json(files_index_path(uid), index)

    return {
        "file_id": file_id, "name": base_name,
        "parsed": result["total"], "added": added,
        "total_in_file": len(existing), "with_answer": has_ans,
        "ans_rate": result["ans_rate"], "parse_method": result["method"],
        "ai_available": result["ai_available"], "message": "Upload thành công"
    }

@app.delete("/files/{file_id}")
def delete_file(file_id: str, x_user_id: str = Header(default="guest")):
    uid = get_uid(x_user_id)
    index = load_json(files_index_path(uid), {})
    if file_id not in index:
        raise HTTPException(404, "Không tìm thấy file")
    q_file = user_dir(uid) / f"{file_id}.json"
    if q_file.exists():
        q_file.unlink()
    del index[file_id]
    save_json(files_index_path(uid), index)
    return {"message": "Đã xóa"}

@app.get("/quiz/start")
def start_quiz(num: int = 10, file_id: str = "",
               x_user_id: str = Header(default="guest")):
    uid = get_uid(x_user_id)
    if file_id:
        all_qs = load_json(user_dir(uid) / f"{file_id}.json", [])
    else:
        index = load_json(files_index_path(uid), {})
        all_qs = []
        for fid in index:
            all_qs.extend(load_json(user_dir(uid) / f"{fid}.json", []))

    valid = [q for q in all_qs if len(q.get("choices", [])) >= 2]
    if not valid:
        raise HTTPException(404, "Không có câu hỏi")

    selected = random.sample(valid, min(num, len(valid)))
    sid = f"s{random.randint(100000, 999999)}"
    quiz = []
    for i, q in enumerate(selected):
        choices = q["choices"].copy()
        random.shuffle(choices)
        correct_text = next((c["text"] for c in q["choices"]
                             if c["label"] == q.get("answer", "")), None)
        new_correct = next((c["label"] for c in choices
                            if c["text"] == correct_text), "") if correct_text else ""
        quiz.append({"id": i, "question": q["question"],
                     "choices": choices, "_correct": new_correct})
    SESSIONS[sid] = quiz
    return {
        "session_id": sid, "total": len(quiz), "file_id": file_id,
        "questions": [{"id": q["id"], "question": q["question"],
                       "choices": q["choices"]} for q in quiz]
    }

class SubmitBody(BaseModel):
    session_id: str
    answers: dict
    time_taken: int = 0
    file_id: str = ""

@app.post("/quiz/submit")
def submit_quiz(body: SubmitBody, x_user_id: str = Header(default="guest")):
    uid = get_uid(x_user_id)
    quiz = SESSIONS.get(body.session_id)
    if not quiz:
        raise HTTPException(404, "Session không tồn tại")
    details = []
    for q in quiz:
        user = body.answers.get(str(q["id"]), "")
        ok = bool(user and user == q["_correct"])
        details.append({"id": q["id"], "question": q["question"],
                        "user": user, "correct": q["_correct"],
                        "ok": ok, "choices": q["choices"]})
    score = sum(1 for d in details if d["ok"])
    pct = round(score / len(quiz) * 100)
    history = load_json(history_path(uid), [])
    history.append({
        "id": len(history)+1,
        "date": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "score": score, "total": len(quiz), "percent": pct,
        "time_taken": body.time_taken, "file_id": body.file_id or "all",
        "wrong_questions": [d["question"][:60] for d in details if not d["ok"]][:5]
    })
    save_json(history_path(uid), history)
    del SESSIONS[body.session_id]
    return {"score": score, "total": len(quiz), "percent": pct, "details": details}

@app.get("/history")
def get_history(x_user_id: str = Header(default="guest")):
    return load_json(history_path(get_uid(x_user_id)), [])

@app.delete("/history/clear")
def clear_history(x_user_id: str = Header(default="guest")):
    save_json(history_path(get_uid(x_user_id)), [])
    return {"message": "Đã xóa lịch sử"}

# ── Debug endpoint ──
@app.post("/debug/parse")
async def debug_parse(file: UploadFile = File(...)):
    """Test parse file — gọi: POST http://127.0.0.1:8000/debug/parse"""
    content = await file.read()
    fname = file.filename or "file"
    try:
        if fname.lower().endswith(".pdf"):
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            page = doc[0]
            hl_rects = get_highlight_rects(page)
            spans = get_all_spans(page)
            table_qs = None
            try:
                table_qs = parse_pdf_table(page, hl_rects)
            except Exception as e:
                table_qs = f"ERROR: {e}"
            inline_qs = parse_pdf_inline(page, hl_rects)
            all_qs = parse_pdf(content)
            return {
                "page_count": len(doc),
                "highlight_rects": len(hl_rects),
                "highlight_colors": [
                    {"fill": list(d.get("fill", [])), "rect": list(d["rect"])}
                    for d in page.get_drawings()
                    if is_highlight(d.get("fill"))
                ][:10],
                "total_spans": len(spans),
                "table_questions": len(table_qs) if isinstance(table_qs, list) else table_qs,
                "inline_questions": len(inline_qs),
                "total_parsed": len(all_qs),
                "with_answer": sum(1 for q in all_qs if q.get("answer")),
                "sample_questions": all_qs[:3]
            }
        else:
            qs = parse_word(content)
            return {"total": len(qs),
                    "with_answer": sum(1 for q in qs if q.get("answer")),
                    "sample": qs[:3]}
    except Exception as e:
        return {"error": str(e), "traceback": traceback.format_exc()}